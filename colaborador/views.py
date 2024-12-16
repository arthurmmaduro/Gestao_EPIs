from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from django.conf import settings
from django.contrib import messages
import pandas as pd
from docx import Document
from datetime import datetime 
import os
from colaborador.models import Colaborador
from colaborador.forms import ColaboradorForm
from autenticacao.models import CustomUser

# Create your views here.

def lista_colaborador(request):
    colaboradores = Colaborador.objects.all().order_by('nome')
    return render(request, 'colaborador/lista_colaborador.html', {'colaboradores': colaboradores})


def detalhes_colaborador(request, colaborador_id):
    colaborador = get_object_or_404(Colaborador, pk=colaborador_id)
    return render(request, 'colaborador/detalhes_colaborador.html', {'colaborador': colaborador})

def cadastrar_colaborador(request):
    if request.method == 'POST':
        form = ColaboradorForm(request.POST, request.FILES)
        if form.is_valid():
            email = form.cleaned_data['email']
            if CustomUser.objects.filter(email=email).exists():
                messages.error(request, 'O email informado já está cadastrado')
                return render (request,'colaborador/cadastrar_colaborador.html', {'form': form})
            usuario = CustomUser.objects.create_user(
                email = email,
                password = 'senha123'
            )

            colaborador = form.save(commit = False)
            colaborador.usuario = usuario
            colaborador.save()

            return redirect('lista_colaborador')
        
    else:
        form = ColaboradorForm()
    return render(request, 'colaborador/cadastrar_colaborador.html', {'form': form})
    
def editar_colaborador(request, colaborador_id):
    colaborador = get_object_or_404(Colaborador, pk=colaborador_id)
    if request.method == 'POST':
        form = ColaboradorForm(request.POST, instance=colaborador)
        if form.is_valid():
            colaborador = form.save(commit=False)
            colaborador.save()
            form.save_m2m()
            return redirect('lista_colaborador')
    else:
        form = ColaboradorForm(instance = colaborador)
        return render(request, 'colaborador/editar_colaborador.html', {'form': form, 'colaborador': colaborador})
    
def exportar_realatorio_colaboradores_epis(request):
    colaboradores = Colaborador.objects.all()
    
    dados_relatorio = []
    
    for colaborador in colaboradores:
        for epi in colaborador.epis.all():
            dados_relatorio.append({
                'Nome do Colaborador': colaborador.nome,
                'Cargo': colaborador.cargo,
                'EPI': epi.nome,
                'Modelo do EPI': epi.modelo,
                'Fabricante': epi.fabricante,
                'CA': epi.numero_ca,
            })
    print(dados_relatorio)

    if not dados_relatorio:
        return HttpResponse('Nenhum dado encontrado para exportação')

    df = pd.DataFrame(dados_relatorio)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    response['Content-Disposition'] = 'attachment; filename="relatorio_colaboradores_epis.xlsx"'

    df.to_excel(response, index=False, engine='openpyxl')

    return response

def gerar_ficha_epi(colaborador, epis):
    
    modelo_path = os.path.join(settings.BASE_DIR, 'documentos', 'Ficha de Entrega de EPIs - formulário.docx')
    doc = Document(modelo_path)

    # Substituir os campos de texto
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "Nome completo:" in cell.text:
                    cell.text = f"Nome completo: {colaborador.nome}"
                elif "Função:" in cell.text:
                    cell.text = f"Função: {colaborador.cargo}"
                    
   # Localizar a tabela de EPIs
    table = None
    for tbl in doc.tables:
        if "Descrição Simplificada" in tbl.cell(0, 0).text:
            table = tbl
            break

    if table is None:
        raise ValueError("Tabela de EPIs não encontrada no modelo.")

    # Preencher a tabela com os EPIs
    for epi in epis:
        row_cells = table.add_row().cells
        row_cells[0].text = epi.nome  # Descrição Simplificada
        row_cells[1].text = "1"  # Quantidade
        row_cells[2].text = epi.numero_ca  # C.A.
        row_cells[3].text = epi.modelo  # Modelo
        row_cells[4].text = epi.fabricante if epi.fabricante else "N/A"  # Fabricante
        row_cells[5].text = epi.validade_ca.strftime("%d/%m/%Y") if epi.validade_ca else "N/A"  # Data de Validade do CA
        row_cells[6].text = datetime.now().strftime("%d/%m/%Y")  # Data da Entrega

    # Preparar o arquivo para download
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Ficha_EPI_{colaborador.nome}.docx'
    doc.save(response)
    return response

def ficha_entrega(request, colaborador_id):
    colaborador = get_object_or_404(Colaborador, id=colaborador_id)
    epis = colaborador.epis.all()  # Ajuste conforme a relação entre os modelos
    return gerar_ficha_epi(colaborador, epis)
